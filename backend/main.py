from __future__ import annotations

import json
import os
import re
from datetime import datetime
from io import BytesIO
from textwrap import wrap
from typing import Any, Dict, List, Literal, Mapping, TypedDict, cast
from uuid import uuid4

from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, Field, field_validator


ModelType = Literal["gemini", "claude", "openai"]


class AnalysisState(TypedDict, total=False):
    project_name: str
    requirement_text: str
    rag_context: str
    api_key: str
    model_type: ModelType
    acceptance_criteria: List[str]
    gherkin_scenarios: List[str]
    story_points: int
    qa_feedback: str
    qa_status: Literal["passed", "needs_review"]


class AnalyzeRequest(BaseModel):
    model_config = ConfigDict(protected_namespaces=())

    session_id: str = Field(default_factory=lambda: str(uuid4()))
    project_name: str
    requirement_text: str
    rag_collection: str = "default"
    api_key: str
    model_type: ModelType = "gemini"

    @field_validator("api_key")
    @classmethod
    def _api_key_nonempty(cls, v: str) -> str:
        s = (v or "").strip()
        if not s:
            raise ValueError("api_key bos olamaz.")
        return s


class AnalysisResponse(BaseModel):
    session_id: str
    analysis_id: str
    project_name: str
    acceptance_criteria: List[str]
    gherkin_scenarios: List[str]
    story_points: int
    qa_feedback: str
    qa_status: str
    rag_context_used: str
    generated_at: str


class GithubExportRequest(BaseModel):
    analysis_id: str
    repository: str
    token: str
    file_path: str = "exports/analysis.md"
    branch: str = "main"
    commit_message: str = "Add AI-RA analysis export"


class JiraExportRequest(BaseModel):
    analysis_id: str
    jira_base_url: str
    email: str
    api_token: str
    project_key: str
    issue_type: str = "Task"


class RegisterAnalysisRequest(BaseModel):
    session_id: str
    project_name: str
    requirement_text: str
    acceptance_criteria: List[str]
    gherkin_scenarios: List[str]
    story_points: int
    qa_feedback: str
    qa_status: Literal["passed", "needs_review"]
    rag_context_used: str = ""


app = FastAPI(title="AI-RA Backend", version="2.0.0", description="Stateless AI-Driven Requirements Analyst API")

_origins_raw = os.environ.get("CORS_ALLOW_ORIGINS", "*").strip()
_cors_origins = [o.strip() for o in _origins_raw.split(",") if o.strip()] or ["*"]
_cors_allow_credentials = "*" not in _cors_origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=_cors_allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
)

_chroma_client = None
session_store: Dict[str, Dict[str, dict]] = {}

_FIB_STORY_POINTS = frozenset({1, 2, 3, 5, 8, 13})

_compiled_analysis_workflow: Any = None


def get_analysis_workflow() -> Any:
    """LangGraph derlemesi ve LangChain yuklemesi yalnizca ilk analiz isteginde yapilir."""
    global _compiled_analysis_workflow
    if _compiled_analysis_workflow is None:
        _compiled_analysis_workflow = _compile_analysis_workflow()
    return _compiled_analysis_workflow


def _compile_analysis_workflow() -> Any:
    from langgraph.graph import END, StateGraph

    def _build_chat_model(model_type: str, api_key: str) -> Any:
        key = (api_key or "").strip()
        if not key:
            raise ValueError("API anahtari bos olamaz.")
        mt = (model_type or "gemini").strip().lower()
        if mt == "gemini":
            from langchain_google_genai import ChatGoogleGenerativeAI

            return ChatGoogleGenerativeAI(
                model="gemini-2.5-flash",
                google_api_key=key,
                temperature=0.2,
            )
        if mt == "claude":
            from langchain_anthropic import ChatAnthropic

            return ChatAnthropic(
                model_name="claude-3-5-sonnet-20240620",
                api_key=key,
                temperature=0.2,
            )
        if mt == "openai":
            from langchain_openai import ChatOpenAI

            return ChatOpenAI(
                model="gpt-4o",
                api_key=key,
                temperature=0.2,
            )
        raise ValueError(f"Desteklenmeyen model_type: {model_type!r}")

    def analyst_agent(state: AnalysisState) -> AnalysisState:
        llm = _build_chat_model(state.get("model_type", "gemini"), state["api_key"])
        requirement_text = state["requirement_text"].strip()
        rag_context = state.get("rag_context", "")
        prompt = (
            "Sen AI-RA Analist Ajanisin.\n"
            "Tum metin ciktilari TURKCE olmalidir.\n"
            "Yalnizca gecerli JSON don; baska aciklama ekleme.\n"
            "acceptance_criteria: 3-7 elemanli dizi; her eleman TEK STRING olmali "
            "(ornek: \"Given kullanici giris sayfasindadir, When gecerli bilgileri girer, Then erisim saglanir\").\n"
            "gherkin_scenarios: 1-3 elemanli dizi; her eleman TEK STRING (cok satirli Gherkin blogu).\n"
            "story_points: Fibonacci tamsayisi (1, 2, 3, 5, 8 veya 13).\n"
            "Ornek JSON:\n"
            '{"acceptance_criteria":["Given ...","When ...","Then ..."],'
            '"gherkin_scenarios":["Feature: X\\n  Scenario: Y\\n    Given ..."],'
            '"story_points":3}\n'
            "Spesifik, test edilebilir ol; halusinasyon yapma.\n\n"
            f"Proje Adi:\n{state['project_name']}\n\n"
            f"Gereksinim:\n{requirement_text}\n\n"
            f"RAG Baglami:\n{rag_context or 'Ek baglam yok'}\n"
        )
        response = llm.invoke(prompt)
        raw = _assistant_text_content(response.content)
        data = _extract_json_payload(raw)
        normalized = _normalize_analyst_payload(data)
        return normalized

    def qa_agent(state: AnalysisState) -> AnalysisState:
        llm = _build_chat_model(state.get("model_type", "gemini"), state["api_key"])
        prompt = (
            "Sen AI-RA QA Ajanisin.\n"
            "Analist ciktisini tutarlilik, belirsizlik ve halusinasyon riski acisindan degerlendir.\n"
            "Yalnizca JSON don: qa_status (passed|needs_review), qa_feedback (string, TURKCE).\n"
            "qa_feedback mutlaka Turkce ve anlasilir olmali.\n\n"
            f"Gereksinim:\n{state['requirement_text']}\n\n"
            f"BDD Kabul Kriterleri:\n{json.dumps(state.get('acceptance_criteria', []), ensure_ascii=False)}\n\n"
            f"Gherkin Senaryolari:\n{json.dumps(state.get('gherkin_scenarios', []), ensure_ascii=False)}\n\n"
            f"Story Points:\n{state.get('story_points')}\n"
        )
        response = llm.invoke(prompt)
        raw = _assistant_text_content(response.content)
        data = _extract_json_payload(raw)
        return _normalize_qa_payload(data)

    graph = StateGraph(AnalysisState)
    graph.add_node("analyst_agent", analyst_agent)
    graph.add_node("qa_agent", qa_agent)
    graph.set_entry_point("analyst_agent")
    graph.add_edge("analyst_agent", "qa_agent")
    graph.add_edge("qa_agent", END)
    return graph.compile()


def _get_chroma_client():
    global _chroma_client
    if _chroma_client is None:
        try:
            import chromadb
        except ImportError as exc:
            raise HTTPException(
                status_code=503,
                detail=(
                    "RAG (ChromaDB) bu ortamda kurulu degil. "
                    "Lokal icin: pip install -r requirements-rag.txt"
                ),
            ) from exc

        persist_dir = os.environ.get("CHROMA_PERSIST_DIR", "").strip()
        if persist_dir:
            _chroma_client = chromadb.PersistentClient(path=persist_dir)
        else:
            _chroma_client = chromadb.Client()
    return _chroma_client


def _fetch_rag_context(collection_name: str, query_text: str) -> str:
    try:
        collection = _get_chroma_client().get_or_create_collection(name=collection_name)
    except HTTPException:
        return ""
    if collection.count() == 0:
        return ""
    query = collection.query(query_texts=[query_text], n_results=3)
    docs = query.get("documents", [[]])[0]
    return "\n".join(docs).strip()


def _build_markdown(analysis: dict) -> str:
    qa_label = "Onaylandi" if analysis.get("qa_status") == "passed" else "Inceleme gerekli"
    lines = [
        f"# AI-RA Analizi — {analysis['project_name']}",
        "",
        f"- Analiz ID: {analysis['analysis_id']}",
        f"- Oturum ID: {analysis['session_id']}",
        f"- Olusturulma: {analysis['generated_at']}",
        f"- Story Point: {analysis['story_points']}",
        f"- QA Durumu: {qa_label}",
        "",
        "## Gereksinim",
        analysis["requirement_text"],
        "",
        "## BDD Kabul Kriterleri",
    ]
    lines.extend([f"- {item}" for item in analysis["acceptance_criteria"]])
    lines.append("")
    lines.append("## Gherkin Senaryolari")
    for scenario in analysis["gherkin_scenarios"]:
        lines.append("```gherkin")
        lines.append(scenario)
        lines.append("```")
        lines.append("")
    lines.append("## QA Geri Bildirimi")
    lines.append(analysis["qa_feedback"])
    return "\n".join(lines).strip()


def _extract_json_payload(text: str) -> dict:
    cleaned = text.strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    match = re.search(r"```json\s*(\{.*?\})\s*```", cleaned, flags=re.DOTALL)
    if not match:
        match = re.search(r"(\{.*\})", cleaned, flags=re.DOTALL)
    if not match:
        raise ValueError("Model JSON ciktisi ayrisamadi.")
    return json.loads(match.group(1))


def _coerce_text_item(item) -> str:
    """Model bazen string yerine nesne/liste dondurur; BDD maddesine cevir."""
    if item is None:
        return ""
    if isinstance(item, str):
        return item.strip()
    if isinstance(item, (int, float, bool)):
        return str(item).strip()
    if isinstance(item, dict):
        given = item.get("given") or item.get("Given") or item.get("GIVEN")
        when = item.get("when") or item.get("When") or item.get("WHEN")
        then = item.get("then") or item.get("Then") or item.get("THEN")
        if any((given, when, then)):
            parts: List[str] = []

            def _bdd_part(label: str, val) -> str:
                text = str(val).strip()
                if not text:
                    return ""
                low = text.lower()
                if low.startswith(label.lower()):
                    return text
                return f"{label} {text}"

            for label, val in (("Given", given), ("When", when), ("Then", then)):
                part = _bdd_part(label, val)
                if part:
                    parts.append(part)
            return ", ".join(parts)
        for key in ("text", "criterion", "description", "content", "bdd", "acceptance_criteria"):
            if key in item and item[key]:
                return _coerce_text_item(item[key])
        str_vals = [str(v).strip() for v in item.values() if v is not None and str(v).strip()]
        return " — ".join(str_vals)
    if isinstance(item, (list, tuple)):
        parts = [_coerce_text_item(x) for x in item]
        return ", ".join(p for p in parts if p)
    return str(item).strip()


def _coerce_string_list(raw, *, field_name: str, min_items: int, max_items: int) -> List[str]:
    items = raw
    if isinstance(items, str):
        items = [line.strip(" •-\t") for line in items.splitlines() if line.strip()]
    if not isinstance(items, list) or not items:
        raise ValueError(f"{field_name} bos veya gecersiz.")
    coerced = [_coerce_text_item(x) for x in items]
    coerced = [x for x in coerced if x]
    if len(coerced) < min_items:
        raise ValueError(
            f"{field_name} yeterli dolu madde icermiyor (en az {min_items}, alinan {len(coerced)})."
        )
    if len(coerced) > max_items:
        coerced = coerced[:max_items]
    return coerced


def _normalize_analyst_payload(data: dict) -> dict:
    if not isinstance(data, dict):
        raise ValueError("Model ciktisi bir JSON nesnesi olmalidir.")
    ac = _coerce_string_list(
        data.get("acceptance_criteria"),
        field_name="acceptance_criteria",
        min_items=3,
        max_items=7,
    )
    gs = _coerce_string_list(
        data.get("gherkin_scenarios"),
        field_name="gherkin_scenarios",
        min_items=1,
        max_items=3,
    )
    sp = data.get("story_points")
    try:
        sp_int = int(sp)
    except (TypeError, ValueError) as exc:
        raise ValueError("story_points tamsayi olmalidir.") from exc
    if sp_int not in _FIB_STORY_POINTS:
        raise ValueError("story_points yalnizca 1, 2, 3, 5, 8 veya 13 olabilir.")
    return {
        "acceptance_criteria": ac,
        "gherkin_scenarios": gs,
        "story_points": sp_int,
    }


def _assistant_text_content(content) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        chunks: List[str] = []
        for part in content:
            if isinstance(part, str):
                chunks.append(part)
            elif isinstance(part, dict):
                text = part.get("text")
                if isinstance(text, str):
                    chunks.append(text)
            else:
                chunks.append(str(part))
        return "".join(chunks)
    return str(content)


def _normalize_qa_payload(data: dict) -> dict:
    if not isinstance(data, dict):
        raise ValueError("QA model ciktisi bir JSON nesnesi olmalidir.")
    status = data.get("qa_status")
    feedback = data.get("qa_feedback")
    if status not in {"passed", "needs_review"}:
        raise ValueError("qa_status yalnizca passed veya needs_review olabilir.")
    if not isinstance(feedback, str) or not feedback.strip():
        raise ValueError("qa_feedback bos olamaz.")
    return {"qa_status": status, "qa_feedback": feedback.strip()}


@app.get("/", tags=["system"])
def root_health() -> dict:
    return {"status": "ok"}


@app.get("/health", tags=["system"])
def health_check() -> dict:
    return {"status": "ok"}


@app.post("/rag/upload", tags=["rag"])
async def upload_rag_pdf(
    session_id: str,
    collection_name: str = "default",
    file: UploadFile = File(...),
) -> dict:
    from pypdf import PdfReader

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Sadece PDF yuklenebilir.")
    reader = PdfReader(BytesIO(await file.read()))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n".join([page for page in pages if page])
    if not text:
        raise HTTPException(status_code=400, detail="PDF icinde okunabilir metin bulunamadi.")

    chunks = [text[i : i + 900] for i in range(0, len(text), 900)]
    collection = _get_chroma_client().get_or_create_collection(name=collection_name)
    ids = [f"{session_id}-{uuid4()}" for _ in chunks]
    metadatas = [{"session_id": session_id, "source": file.filename}] * len(chunks)
    collection.add(ids=ids, documents=chunks, metadatas=metadatas)
    return {"status": "indexed", "chunks_added": len(chunks), "collection_name": collection_name}


@app.post("/analyze", response_model=AnalysisResponse, tags=["analysis"])
def analyze_requirement(http_request: Request, body: AnalyzeRequest) -> AnalysisResponse:
    rag_context = _fetch_rag_context(body.rag_collection, body.requirement_text)
    api_key = body.api_key
    hdr_model = (http_request.headers.get("x-model-type") or "").strip().lower()
    model_type: ModelType = body.model_type
    if hdr_model in ("gemini", "claude", "openai"):
        model_type = cast(ModelType, hdr_model)
    try:
        result = get_analysis_workflow().invoke(
            {
                "project_name": body.project_name,
                "requirement_text": body.requirement_text,
                "rag_context": rag_context,
                "api_key": api_key,
                "model_type": model_type,
            }
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"LLM analizi basarisiz: {exc}") from exc
    analysis_id = str(uuid4())
    payload = {
        "session_id": body.session_id,
        "analysis_id": analysis_id,
        "project_name": body.project_name,
        "requirement_text": body.requirement_text,
        "acceptance_criteria": result["acceptance_criteria"],
        "gherkin_scenarios": result["gherkin_scenarios"],
        "story_points": result["story_points"],
        "qa_feedback": result["qa_feedback"],
        "qa_status": result["qa_status"],
        "rag_context_used": rag_context,
        "generated_at": datetime.utcnow().isoformat() + "Z",
    }
    session_store.setdefault(body.session_id, {})[analysis_id] = payload
    return AnalysisResponse(**payload)


@app.get("/analysis/{session_id}/{analysis_id}", tags=["analysis"])
def get_analysis(session_id: str, analysis_id: str) -> dict:
    analysis = session_store.get(session_id, {}).get(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analiz bulunamadi.")
    return analysis


@app.post("/analysis/register", tags=["analysis"])
def register_analysis(body: RegisterAnalysisRequest) -> dict:
    """Django veya diger istemcilerin export oncesi analiz kaydi olusturmasi icin."""
    analysis_id = str(uuid4())
    payload = {
        "session_id": body.session_id,
        "analysis_id": analysis_id,
        "project_name": body.project_name,
        "requirement_text": body.requirement_text,
        "acceptance_criteria": body.acceptance_criteria,
        "gherkin_scenarios": body.gherkin_scenarios,
        "story_points": body.story_points,
        "qa_feedback": body.qa_feedback,
        "qa_status": body.qa_status,
        "rag_context_used": body.rag_context_used,
        "generated_at": datetime.utcnow().isoformat() + "Z",
    }
    session_store.setdefault(body.session_id, {})[analysis_id] = payload
    return {"session_id": body.session_id, "analysis_id": analysis_id}


@app.get("/export/pdf/{session_id}/{analysis_id}", tags=["export"])
def export_pdf(session_id: str, analysis_id: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    analysis = session_store.get(session_id, {}).get(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analiz bulunamadi.")
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    text_obj = pdf.beginText(40, 800)
    for line in _build_markdown(analysis).splitlines():
        for wrapped in wrap(line, 100) or [""]:
            text_obj.textLine(wrapped)
    pdf.drawText(text_obj)
    pdf.showPage()
    pdf.save()
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=analysis-{analysis_id}.pdf"},
    )


@app.get("/export/word/{session_id}/{analysis_id}", tags=["export"])
def export_word(session_id: str, analysis_id: str):
    from docx import Document

    analysis = session_store.get(session_id, {}).get(analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analiz bulunamadi.")
    doc = Document()
    doc.add_heading(f"AI-RA Analysis - {analysis['project_name']}", 1)
    doc.add_paragraph(f"Analysis ID: {analysis_id}")
    doc.add_paragraph(f"Story Points: {analysis['story_points']}")
    doc.add_heading("BDD Acceptance Criteria", level=2)
    for item in analysis["acceptance_criteria"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Gherkin Scenarios", level=2)
    for scenario in analysis["gherkin_scenarios"]:
        doc.add_paragraph(scenario)
    doc.add_heading("QA Feedback", level=2)
    doc.add_paragraph(analysis["qa_feedback"])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=analysis-{analysis_id}.docx"},
    )


def _merge_github_export(body: GithubExportRequest, headers: Mapping[str, str]) -> GithubExportRequest:
    data = body.model_dump()
    if not str(data.get("token", "")).strip():
        data["token"] = (headers.get("x-github-token") or "").strip()
    if not str(data.get("repository", "")).strip():
        data["repository"] = (headers.get("x-github-repository") or "").strip()
    return GithubExportRequest(**data)


def _merge_jira_export(body: JiraExportRequest, headers: Mapping[str, str]) -> JiraExportRequest:
    data = body.model_dump()
    pairs = [
        ("jira_base_url", "x-jira-base-url"),
        ("email", "x-jira-email"),
        ("api_token", "x-jira-api-token"),
        ("project_key", "x-jira-project-key"),
    ]
    for field, hdr in pairs:
        if not str(data.get(field, "")).strip():
            data[field] = (headers.get(hdr) or "").strip()
    return JiraExportRequest(**data)


@app.post("/export/github/{session_id}", tags=["export"])
def export_to_github(session_id: str, http_request: Request, body: GithubExportRequest) -> dict:
    import requests

    export_req = _merge_github_export(body, http_request.headers)
    if not export_req.repository.strip() or not export_req.token.strip():
        raise HTTPException(
            status_code=400,
            detail="GitHub export icin repository ve token zorunludur.",
        )
    analysis = session_store.get(session_id, {}).get(export_req.analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analiz bulunamadi.")
    api_url = f"https://api.github.com/repos/{export_req.repository}/contents/{export_req.file_path}"
    content = _build_markdown(analysis).encode("utf-8")
    payload = {
        "message": export_req.commit_message,
        "content": __import__("base64").b64encode(content).decode("utf-8"),
        "branch": export_req.branch,
    }
    headers = {"Authorization": f"Bearer {export_req.token}", "Accept": "application/vnd.github+json"}
    response = requests.put(api_url, json=payload, headers=headers, timeout=15)
    if response.status_code not in (200, 201):
        raise HTTPException(status_code=400, detail=f"GitHub export hatasi: {response.text}")
    return {"status": "ok", "target": export_req.repository, "path": export_req.file_path}


@app.post("/export/jira/{session_id}", tags=["export"])
def export_to_jira(session_id: str, http_request: Request, body: JiraExportRequest) -> dict:
    import requests

    export_req = _merge_jira_export(body, http_request.headers)
    if not all(
        [
            export_req.jira_base_url.strip(),
            export_req.email.strip(),
            export_req.api_token.strip(),
            export_req.project_key.strip(),
        ]
    ):
        raise HTTPException(
            status_code=400,
            detail="Jira export icin base URL, email, token ve project_key zorunludur.",
        )
    analysis = session_store.get(session_id, {}).get(export_req.analysis_id)
    if not analysis:
        raise HTTPException(status_code=404, detail="Analiz bulunamadi.")
    issue_url = f"{export_req.jira_base_url.rstrip('/')}/rest/api/3/issue"
    markdown = _build_markdown(analysis)
    payload = {
        "fields": {
            "project": {"key": export_req.project_key},
            "summary": f"AI-RA Analysis {export_req.analysis_id}",
            "issuetype": {"name": export_req.issue_type},
            "description": {
                "type": "doc",
                "version": 1,
                "content": [
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": markdown[:32767]}],
                    }
                ],
            },
        }
    }
    response = requests.post(issue_url, json=payload, auth=(export_req.email, export_req.api_token), timeout=15)
    if response.status_code not in (200, 201):
        raise HTTPException(status_code=400, detail=f"Jira export hatasi: {response.text}")
    data = response.json()
    return {"status": "ok", "jira_issue_key": data.get("key")}
